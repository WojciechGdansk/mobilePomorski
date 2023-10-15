from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from form_data.form_info import RegionalLeague, FourthLeague
from form_data.match_info import MatchInfo
from word_creator.word_table_data import TableData


class CreateWord:
    doc = Document()
    page_size = (21, 29.7)
    margins = (2.1, 2.1, 2.1, 1.7)

    def __init__(self, match_data_object: MatchInfo, stadium_facilities: [FourthLeague, RegionalLeague], user_name):
        self.match_data = match_data_object
        self.stadium_facilities = stadium_facilities
        self.user_name = user_name

    def set_margin(self):
        section = self.doc.sections[0]
        section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = map(Cm, self.margins)

    def set_page_size(self):
        section = self.doc.sections[0]
        section.page_width, section.page_height = map(Cm, self.page_size)

    def save_document(self, document_name="Załącznik do sprawozdania - MW.docx"):
        self.doc.save(f'{document_name}')

    def set_normal_style(self):
        default_font = self.doc.styles['Normal'].font
        default_font.name = "Cambria"
        default_font.size = Pt(12)

    def section_before_table(self):
        empty_line = self.doc.add_paragraph(' ', style="No Spacing")
        run = empty_line.runs[0]
        run.font.size = Pt(12)
        paragraph = self.doc.add_paragraph('ZAŁĄCZNIK DO SPRAWOZDANIA SĘDZIEGO LUB OBSERWATORA SĘDZIÓW', style="Normal")
        # Customize paragraph alignment
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Access the Run object within the paragraph
        run = paragraph.runs[0]
        # Change the font size
        run.font.size = Pt(14)
        # Make the text bold
        run.bold = True

    def add_custom_style(self):
        styles = self.doc.styles
        style = styles.add_style('Small space', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["No Spacing"]
        font = style.font
        font.size = Pt(4)

    def draw_table(self):
        table = self.doc.add_table(rows=1, cols=4, style="Table Grid")
        for row in table.rows:
            cell = row.cells[0]
            cell.width = Cm(4.48)
            paragraph = cell.paragraphs[0]
            paragraph.add_run("Zawody ")
            italic_run = paragraph.add_run("(liga/klasa):")
            italic_run.italic = True

            row.cells[1].width = Cm(4.25)
            row.cells[1].text = self.match_data.league
            row.cells[2].text = "Data"
            row.cells[2].width = Cm(4.31)
            row.cells[3].text = self.match_data.date
            row.cells[3].width = Cm(4.04)

        table = self.doc.add_table(rows=1, cols=4, style='Table Grid')

        for row in table.rows:
            row.cells[0].text = "Gospodarze"
            row.cells[1].text = self.match_data.home_team
            row.cells[0].width = Cm(2.69)
            row.cells[2].text = "Goście"
            row.cells[3].text = self.match_data.away_team

    def draw_empty_row_in_table(self, size_of_font):
        table = self.doc.add_table(rows=1, cols=1, style="Table Grid")
        for row in table.rows:
            cell = row.cells[0]
            paragraph = cell.paragraphs[0]
            empty_line = paragraph.add_run(' ')
            empty_line.font.size = Pt(size_of_font)

    def draw_main_table(self):
        for index, key in enumerate(TableData.dane):
            if key == "other_remarks":
                self.doc.add_paragraph('', style="Small space")
                table = self.doc.add_table(rows=1, cols=2, style="Table Grid")
                row = table.rows[0]
                row.height = Cm(9.02)
                row.cells[0].width = Cm(0.94)
                row.cells[0].vertical_alignment = WD_TAB_ALIGNMENT.CENTER
                row.cells[0].text = str(index + 1)
                row.cells[1].width = Cm(16.15)
                row.cells[1].text = TableData.dane[key]
                other_remarks = self.stadium_facilities.other_remarks
                row.cells[1].add_paragraph(f'{other_remarks if type(other_remarks) is not bool else ""}')
            else:
                table = self.doc.add_table(rows=1, cols=4, style="Table Grid")
                for row in table.rows:
                    row.cells[0].text = str(index + 1)
                    row.cells[1].text = TableData.dane[key]
                    selected = None
                    try:
                        selected = self.stadium_facilities.__getattribute__(key)
                    except AttributeError:
                        pass
                    if selected:
                        paragraph = row.cells[2].paragraphs[0]
                        text_to_format = paragraph.add_run("TAK")
                        text_to_format.underline = True
                        text_to_format.bold = True
                        row.cells[3].text = "NIE"
                    if not selected:
                        paragraph = row.cells[3].paragraphs[0]
                        text_to_format = paragraph.add_run("NIE")
                        text_to_format.underline = True
                        text_to_format.bold = True
                        row.cells[2].text = "TAK"
                    if selected is None:
                        row.cells[2].text = "TAK"
                        row.cells[3].text = "NIE"

                    row.cells[0].width = Cm(0.94)
                    row.cells[1].width = Cm(11.65)
                    row.cells[2].width = Cm(2.25)
                    row.cells[3].width = Cm(2.25)
                self.draw_empty_row_in_table(8)

    def last_row(self):
        self.doc.add_paragraph('', style="Small space")
        table = self.doc.add_table(rows=1, cols=2, style="Table Grid")
        row = table.rows[0]
        row.cells[0].text = "Imię i nazwisko sporządzającego "
        row.cells[1].text = self.user_name

    def create_document(self):
        self.set_page_size()
        self.set_margin()
        self.set_normal_style()
        self.section_before_table()
        self.add_custom_style()
        self.draw_table()
        self.doc.add_paragraph(" ", style='No Spacing')
        self.draw_main_table()
        self.last_row()
        self.save_document()
